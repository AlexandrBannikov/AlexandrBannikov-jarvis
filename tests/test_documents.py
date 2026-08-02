"""Security and lifecycle coverage for ephemeral documents."""

from datetime import datetime, timedelta, timezone
from pathlib import Path
import os
import sqlite3

import pytest

from app.documents.chunking import DocumentChunker, normalize
from app.documents.formatter import redact, provenance
from app.documents.models import DocumentSession, ExtractedPart, Provenance
from app.documents.service import DocumentService
from app.documents.storage import DocumentSessionStorage
from app.documents.validators import DocumentValidator, ValidationError
from app.documents.extractors import ExtractionError, SpreadsheetExtractor, TextExtractor


@pytest.mark.parametrize(("name","mime","kind"),[
 ("a.txt","text/plain","txt"),("a.md","text/markdown","md"),("a.pdf","application/pdf","pdf"),
 ("a.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document","docx"),
 ("a.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet","xlsx"),
 ("a.csv","text/csv","csv"),("a.json","application/json","json"),("a.jpg","image/jpeg","image"),
 ("a.jpeg","image/jpeg","image"),("a.png","image/png","image"),("a.webp","image/webp","image"),
 ("a.MD","text/plain","md"),("данные.txt","text/plain","txt"),("report.csv","application/csv","csv"),
 ("data.json","text/plain","json"),
])
def test_supported_metadata(name,mime,kind):
    assert DocumentValidator(100).validate_metadata(name,mime,10).document_type==kind


@pytest.mark.parametrize(("name","mime","code"),[
 ("x.exe","application/pdf","UNSUPPORTED_FORMAT"),("x.zip","application/zip","UNSUPPORTED_FORMAT"),
 ("x.sh","text/plain","UNSUPPORTED_FORMAT"),("x.pdf","text/plain","MIME_MISMATCH"),
 ("x.png","image/jpeg","MIME_MISMATCH"),("x.docx","application/zip","MIME_MISMATCH"),
 ("x.xlsm","application/vnd.ms-excel.sheet.macroEnabled.12","UNSUPPORTED_FORMAT"),
 ("x.docm","application/vnd.ms-word.document.macroEnabled.12","UNSUPPORTED_FORMAT"),
 ("x.mp4","video/mp4","UNSUPPORTED_FORMAT"),("x.mp3","audio/mpeg","UNSUPPORTED_FORMAT"),
])
def test_rejected_metadata(name,mime,code):
    with pytest.raises(ValidationError) as exc:DocumentValidator(100).validate_metadata(name,mime,10)
    assert exc.value.code==code


@pytest.mark.parametrize(("unsafe","safe"),[
 ("../../etc/passwd.pdf","passwd.pdf"),("..\\..\\x.txt","x.txt"),("/tmp/x.csv","x.csv"),
 ("a?b.json","a_b.json"),(" a.pdf ","a.pdf"),("...","document"),(None,"document"),
 ("a\x00b.txt","a_b.txt"),("dir/sub/file.md","file.md"),("тест✅.txt","тест_.txt"),
])
def test_safe_filename(unsafe,safe):assert DocumentValidator.safe_filename(unsafe)==safe


@pytest.mark.parametrize(("suffix","body","valid"),[
 (".pdf",b"%PDF-1.7",True),(".pdf",b"MZ",False),(".png",b"\x89PNG\r\n\x1a\n",True),
 (".png",b"GIF89a",False),(".jpg",b"\xff\xd8\xffx",True),(".jpg",b"PK\x03\x04",False),
 (".docx",b"PK\x03\x04",True),(".docx",b"MZ",False),(".xlsx",b"PK\x03\x04",True),
 (".xlsx",b"#!/bin/sh",False),(".webp",b"RIFF1234WEBP",True),(".webp",b"RIFF1234NOPE",False),
])
def test_magic_bytes(tmp_path,suffix,body,valid):
    p=tmp_path/("x"+suffix);p.write_bytes(body)
    if valid:DocumentValidator(100).validate_content(p,suffix)
    else:
        with pytest.raises(ValidationError):DocumentValidator(100).validate_content(p,suffix)


@pytest.mark.parametrize("text",[
 "hello world","Привет мир","A-heading exact phrase","one two three","warranty 24 months",
 "sheet finance total","page twelve obligation","alpha beta","JSON key value","CSV row value",
])
def test_chunk_search(text):
    chunker=DocumentChunker(100);chunks=chunker.chunk([ExtractedPart(text,Provenance(page=2))])
    assert chunker.search(chunks,text,page=2,limit=1)[0].text==text


@pytest.mark.parametrize(("secret","marker"),[
 ("sk-"+"abcdefghijklmnop","[СКРЫТО]"),("123456789:"+"ABCDEFGHIJKLMNOPQRSTUVWXYZ","[СКРЫТО]"),
 ("password=secret123","password=[СКРЫТО]"),("PASSWORD: qwerty","PASSWORD: [СКРЫТО]"),
 ("before sk-12345678 after","[СКРЫТО]"),("password = hidden","password = [СКРЫТО]"),
])
def test_redaction(secret,marker):assert marker in redact(secret) and "⚠️" in redact(secret)


@pytest.fixture
def storage(tmp_path):
    value=DocumentSessionStorage(tmp_path/"document_sessions.db",tmp_path/"documents");value.initialize();return value

def make_session(storage,user=1,chat=2,identifier="doc",expired=False,message_id=10):
    now=datetime.now(timezone.utc);path=storage.storage_path/(identifier+".txt");path.write_text("private",encoding="utf-8");os.chmod(path,0o600)
    session=DocumentSession(identifier,user,chat,message_id,"filehash",identifier+".txt",identifier+".txt","text/plain",7,"sha","ready","txt",7,0,now,now-timedelta(hours=1) if expired else now+timedelta(hours=1),now,None,True,path)
    storage.insert(session,DocumentChunker().chunk([ExtractedPart("private warranty",Provenance(section="A"))]));return session

def test_storage_permissions(storage):assert (storage.storage_path.stat().st_mode&0o777)==0o700 and (storage.db_path.stat().st_mode&0o777)==0o600
def test_storage_schema(storage):assert storage.validate_schema()
def test_storage_create_and_active(storage):s=make_session(storage);assert storage.active(1,2).id==s.id
def test_storage_user_isolation(storage):make_session(storage);assert storage.get("doc",99,2) is None
def test_storage_chat_isolation(storage):make_session(storage);assert storage.get("doc",1,99) is None
def test_storage_chunk_isolation(storage):make_session(storage);assert storage.chunks("doc",99,2)==[]
def test_storage_list_owner(storage):make_session(storage);assert len(storage.list(1,2))==1 and storage.list(2,2)==[]
def test_forget_deletes_content(storage):s=make_session(storage);assert storage.forget("doc",1,2) and not s.file_path.exists()
def test_forget_deactivates(storage):make_session(storage);storage.forget("doc",1,2);assert storage.active(1,2) is None
def test_forget_wrong_owner(storage):s=make_session(storage);assert not storage.forget("doc",2,2) and s.file_path.exists()
def test_cleanup_dry_run(storage):s=make_session(storage,expired=True);assert storage.cleanup(False)==1 and s.file_path.exists()
def test_cleanup_apply(storage):s=make_session(storage,expired=True);assert storage.cleanup(True)==1 and not s.file_path.exists()
def test_metrics(storage):make_session(storage);make_session(storage,identifier="old",expired=True,message_id=11);assert storage.metrics()=={"active":1,"expired":1}
def test_duplicate_update_rejected(storage):
    make_session(storage)
    with pytest.raises(sqlite3.IntegrityError):make_session(storage,identifier="other")
def test_same_hash_two_users(storage):make_session(storage);make_session(storage,user=2,identifier="other");assert len(storage.list(2,2))==1

@pytest.mark.parametrize(("page","sheet","expected"),[(3,None,"стр. 3"),(None,"Finance","лист Finance"),(3,"F","стр. 3"),(None,None,"section")])
def test_provenance(page,sheet,expected):
    chunk=DocumentChunker().chunk([ExtractedPart("x",Provenance(page=page,sheet=sheet,section="section"))])[0]
    assert expected in provenance(chunk)

@pytest.mark.parametrize("value",["Test TEST","Привет, привет","one-two three","AA BB CC","123 value"])
def test_normalization(value):assert isinstance(normalize(value),frozenset) and normalize(value)

def test_context_bounded(storage):
    make_session(storage);service=DocumentService(storage,max_context_chars=30,max_chunks_per_request=1)
    assert len(service.context(1,2,"warranty"))<200
def test_context_has_private_policy(storage):make_session(storage);assert "private" in DocumentService(storage).context(1,2,"warranty").lower()
def test_context_missing_owner(storage):make_session(storage);assert DocumentService(storage).context(9,2,"warranty")==""
def test_service_forget_active(storage):make_session(storage);assert DocumentService(storage).forget(1,2)
def test_service_list_owned(storage):make_session(storage);assert len(DocumentService(storage).list_documents(1,2))==1
def test_oversize_before_download():
    with pytest.raises(ValidationError) as exc:DocumentValidator(5).validate_metadata("x.txt","text/plain",6)
    assert exc.value.code=="FILE_TOO_LARGE"

def extractor():return TextExtractor(max_chars=10000,max_pdf_pages=10,max_docx_paragraphs=100)
def test_utf8_extraction(tmp_path):p=tmp_path/"a.txt";p.write_bytes("Привет".encode());assert extractor().extract(p,"txt")[0].text=="Привет"
def test_cp1251_extraction(tmp_path):p=tmp_path/"a.txt";p.write_bytes("Привет".encode("cp1251"));assert extractor().extract(p,"txt")[0].text=="Привет"
def test_json_extraction(tmp_path):p=tmp_path/"a.json";p.write_text('{"a": 1}');assert '"a": 1' in extractor().extract(p,"json")[0].text
def test_malformed_json(tmp_path):
    p=tmp_path/"a.json";p.write_text("{")
    with pytest.raises(ExtractionError):extractor().extract(p,"json")
def test_pdf_without_text(tmp_path):
    from pypdf import PdfWriter
    p=tmp_path/"a.pdf";w=PdfWriter();w.add_blank_page(100,100)
    with p.open("wb") as stream:w.write(stream)
    with pytest.raises(ExtractionError) as exc:extractor().extract(p,"pdf")
    assert exc.value.code=="PDF_NO_TEXT"
def test_encrypted_pdf(tmp_path):
    from pypdf import PdfWriter
    p=tmp_path/"a.pdf";w=PdfWriter();w.add_blank_page(100,100);w.encrypt("secret")
    with p.open("wb") as stream:w.write(stream)
    with pytest.raises(ExtractionError) as exc:extractor().extract(p,"pdf")
    assert exc.value.code=="ENCRYPTED_DOCUMENT"
def test_docx_paragraphs_and_tables(tmp_path):
    from docx import Document
    p=tmp_path/"a.docx";d=Document();d.add_paragraph("Obligation");t=d.add_table(rows=1,cols=2);t.cell(0,0).text="Warranty";t.cell(0,1).text="24";d.save(p)
    text="\n".join(x.text for x in extractor().extract(p,"docx"));assert "Obligation" in text and "Warranty | 24" in text
def test_corrupted_docx(tmp_path):
    p=tmp_path/"a.docx";p.write_bytes(b"PK\x03\x04bad")
    with pytest.raises(ExtractionError):extractor().extract(p,"docx")
def test_xlsx_sheets_values_and_formulas(tmp_path):
    from openpyxl import Workbook
    p=tmp_path/"a.xlsx";w=Workbook();s=w.active;s.title="Finance";s.append(["Total",42]);s.append(["Formula","=1+1"]);w.save(p)
    parts=SpreadsheetExtractor(max_cells=100,max_chars=1000).extract(p)
    assert parts[0].provenance.sheet=="Finance" and "42" in parts[0].text and all("=1+1" not in x.text for x in parts)
def test_xlsx_cell_limit(tmp_path):
    from openpyxl import Workbook
    p=tmp_path/"a.xlsx";w=Workbook();w.active.append([1,2,3]);w.save(p)
    with pytest.raises(ExtractionError) as exc:SpreadsheetExtractor(max_cells=2,max_chars=100).extract(p)
    assert exc.value.code=="TOO_MANY_CELLS"
def test_image_pixel_limit(storage,tmp_path):
    from PIL import Image
    p=tmp_path/"x.png";Image.new("RGB",(20,20)).save(p)
    with pytest.raises(Exception):DocumentService(storage,max_image_pixels=100)._validate_image(p)
def test_repeated_message_lookup(storage):s=make_session(storage);assert storage.by_message(1,2,10).id==s.id
