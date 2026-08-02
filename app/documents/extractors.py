"""Non-executing text and spreadsheet extractors."""

from pathlib import Path
import csv
import io
import json
import zipfile

from .models import ExtractedPart, Provenance


class ExtractionError(ValueError):
    def __init__(self, code: str, message: str): super().__init__(message); self.code=code; self.user_message=message


def _decode(data: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1251", "latin-1"):
        try: return data.decode(encoding)
        except UnicodeDecodeError: pass
    raise ExtractionError("DECODE_ERROR", "Не удалось определить кодировку файла.")


class TextExtractor:
    def __init__(self, *, max_chars: int, max_pdf_pages: int, max_docx_paragraphs: int):
        self.max_chars=max_chars; self.max_pdf_pages=max_pdf_pages; self.max_docx_paragraphs=max_docx_paragraphs

    def extract(self, path: Path, kind: str) -> list[ExtractedPart]:
        try:
            if kind in {"txt", "md", "csv", "json"}: return self._plain(path, kind)
            if kind == "pdf": return self._pdf(path)
            if kind == "docx": return self._docx(path)
        except ExtractionError: raise
        except Exception as error: raise ExtractionError("MALFORMED_DOCUMENT", "Документ повреждён или имеет неподдерживаемую структуру.") from error
        raise ExtractionError("UNSUPPORTED_FORMAT", "Этот формат файла пока не поддерживается.")

    def _plain(self, path: Path, kind: str) -> list[ExtractedPart]:
        text=_decode(path.read_bytes())
        if kind == "json":
            obj=json.loads(text); text=json.dumps(obj, ensure_ascii=False, indent=2)
        if len(text)>self.max_chars: raise ExtractionError("TEXT_TOO_LARGE", "Извлечённый текст превышает допустимый объём.")
        return [ExtractedPart(text, Provenance(section=kind.upper()))]

    def _pdf(self, path: Path) -> list[ExtractedPart]:
        from pypdf import PdfReader
        reader=PdfReader(path, strict=True)
        if reader.is_encrypted: raise ExtractionError("ENCRYPTED_DOCUMENT", "Защищённые паролем документы не поддерживаются.")
        if len(reader.pages)>self.max_pdf_pages: raise ExtractionError("TOO_MANY_PAGES", "В PDF слишком много страниц.")
        parts=[ExtractedPart(page.extract_text() or "", Provenance(page=i)) for i,page in enumerate(reader.pages,1)]
        if not any(p.text.strip() for p in parts): raise ExtractionError("PDF_NO_TEXT", "В документе не обнаружен текстовый слой. OCR пока не поддерживается.")
        self._check_chars(parts); return parts

    def _docx(self, path: Path) -> list[ExtractedPart]:
        if not zipfile.is_zipfile(path): raise ExtractionError("MALFORMED_DOCUMENT", "DOCX повреждён.")
        from docx import Document
        doc=Document(path); parts=[]; count=0
        for p in doc.paragraphs:
            count+=1
            if count>self.max_docx_paragraphs: raise ExtractionError("TOO_MANY_PARAGRAPHS", "В DOCX слишком много абзацев.")
            if p.text.strip(): parts.append(ExtractedPart(p.text, Provenance(section="paragraph")))
        for ti,table in enumerate(doc.tables,1):
            for row in table.rows:
                parts.append(ExtractedPart(" | ".join(c.text for c in row.cells), Provenance(section=f"table {ti}")))
        self._check_chars(parts); return parts

    def _check_chars(self, parts):
        if sum(len(p.text) for p in parts)>self.max_chars: raise ExtractionError("TEXT_TOO_LARGE", "Извлечённый текст превышает допустимый объём.")


class SpreadsheetExtractor:
    def __init__(self, *, max_cells: int, max_chars: int): self.max_cells=max_cells; self.max_chars=max_chars
    def extract(self, path: Path) -> list[ExtractedPart]:
        from openpyxl import load_workbook
        try: workbook=load_workbook(path, read_only=True, data_only=True, keep_links=False)
        except Exception as error: raise ExtractionError("MALFORMED_DOCUMENT", "XLSX повреждён.") from error
        parts=[]; cells=0; chars=0
        try:
            for sheet in workbook.worksheets:
                for row_no,row in enumerate(sheet.iter_rows(values_only=True),1):
                    cells += len(row)
                    if cells>self.max_cells: raise ExtractionError("TOO_MANY_CELLS", "В таблице слишком много ячеек.")
                    values=["" if value is None else str(value) for value in row]
                    if any(values):
                        text=" | ".join(values); chars+=len(text)
                        if chars>self.max_chars: raise ExtractionError("TEXT_TOO_LARGE", "Извлечённый текст превышает допустимый объём.")
                        parts.append(ExtractedPart(text, Provenance(sheet=sheet.title,row_start=row_no,row_end=row_no)))
        finally: workbook.close()
        return parts
